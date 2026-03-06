"""
DOCX Writer — Consumes docx_instructions from TRANSFORM agent.

Two modes:
  1. build_apa_docx(): New APA-specific writer that handles the section-based
     format (title_page, abstract_page, body, references_page) from the
     APA Pipeline prompts. This is the PRIMARY path for APA formatting.

  2. write_formatted_docx(): Legacy generic writer for non-APA journals.
     Handles flat section lists (title, heading, paragraph, reference, etc.).

  3. transform_docx_in_place(): In-place transformation of uploaded DOCX files.
     Preserves figures, tables, and embedded objects.

Based on APA_Pipeline_Complete_Prompts.md §6.
"""
import re
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips

from tools.logger import get_logger
from tools.tool_errors import DocumentWriteError

logger = get_logger(__name__)

_CASE_OPTIONS = {"Title Case", "UPPERCASE", "Sentence case", "lowercase"}


# ═══════════════════════════════════════════════════════════════════════════════
# APA-SPECIFIC DOCX BUILDER (from APA Pipeline Prompts §6)
# ═══════════════════════════════════════════════════════════════════════════════

def build_apa_docx(transform_output: dict, output_path: str) -> str:
    """Main entry point: takes TRANSFORM agent JSON → produces APA DOCX file."""

    instructions = transform_output.get("docx_instructions", transform_output)
    sections_data = instructions.get("sections", [])

    if not sections_data:
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    doc = Document()

    # ── 1. SET DOCUMENT-LEVEL DEFAULTS ──
    style = doc.styles['Normal']
    font = style.font
    font.name = instructions.get("font", "Times New Roman")
    # Support both field names: font_size_halfpoints (MD prompt) and font_size (legacy)
    raw_size = instructions.get("font_size_halfpoints", instructions.get("font_size", 24))
    font.size = Pt(raw_size / 2) if raw_size > 13 else Pt(raw_size)

    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(0)
    paragraph_format.line_spacing = 2.0
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Set default font in XML
    rpr = doc.styles['Normal'].element.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rFonts.set(qn('w:cs'), 'Times New Roman')
    rpr.append(rFonts)

    # ── 2. CONFIGURE HEADING STYLES ──
    _configure_heading_styles(doc)

    # ── 3. PROCESS EACH SECTION ──
    first_section = True
    for section_data in sections_data:
        section_type = section_data.get("type", "body")

        if section_type == "title_page":
            _write_title_page(doc, section_data, instructions, first_section)
        elif section_type == "abstract_page":
            _write_abstract_page(doc, section_data, instructions)
        elif section_type == "body":
            _write_body(doc, section_data, instructions)
        elif section_type == "references_page":
            _write_references_page(doc, section_data, instructions)
        else:
            _write_body(doc, section_data, instructions)

        first_section = False

    # ── 4. SET PAGE SIZE & MARGINS ON ALL SECTIONS ──
    page_size = instructions.get("page_size", {})
    margins = instructions.get("margins", {})
    for section in doc.sections:
        section.page_width = Twips(page_size.get("width", 12240))
        section.page_height = Twips(page_size.get("height", 15840))
        section.top_margin = Twips(margins.get("top", 1440))
        section.bottom_margin = Twips(margins.get("bottom", 1440))
        section.left_margin = Twips(margins.get("left", 1440))
        section.right_margin = Twips(margins.get("right", 1440))

        _add_page_number_header(section)

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("[DOCX_APA] Written | file=%s | sections=%d", out.name, len(sections_data))
    return str(out)


def _configure_heading_styles(doc):
    """Set up APA heading styles in the document."""
    h1 = doc.styles['Heading 1']
    h1.font.name = 'Times New Roman'
    h1.font.size = Pt(12)
    h1.font.bold = True
    h1.font.italic = False
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(0)
    h1.paragraph_format.line_spacing = 2.0
    h1.paragraph_format.first_line_indent = Inches(0)

    h2 = doc.styles['Heading 2']
    h2.font.name = 'Times New Roman'
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.italic = False
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(0)
    h2.paragraph_format.space_after = Pt(0)
    h2.paragraph_format.line_spacing = 2.0
    h2.paragraph_format.first_line_indent = Inches(0)

    h3 = doc.styles['Heading 3']
    h3.font.name = 'Times New Roman'
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.italic = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(0)
    h3.paragraph_format.space_after = Pt(0)
    h3.paragraph_format.line_spacing = 2.0
    h3.paragraph_format.first_line_indent = Inches(0.5)


def _write_title_page(doc, section_data, instructions, is_first):
    """Write the APA title page."""
    if not is_first:
        doc.add_section()

    # Default 3 blank lines if no explicit spacing element comes first
    elements = section_data.get("elements", [])
    has_leading_spacing = elements and elements[0].get("type") == "spacing"
    if not has_leading_spacing:
        for _ in range(3):
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

    for element in elements:
        etype = element.get("type", "")

        if etype == "spacing":
            # Blank lines from the MD prompt schema
            blank_lines = element.get("blank_lines", 1)
            for _ in range(blank_lines):
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 2.0
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)

        elif etype == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(element.get("text", ""))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype in ("authors", "affiliation"):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(element.get("text", ""))
            run.bold = False
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)


def _write_abstract_page(doc, section_data, instructions):
    """Write the APA abstract page."""
    doc.add_section()

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "abstract_label":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", "Abstract"))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "abstract_body":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            _add_text_with_italics(p, element.get("text", ""))

        elif etype == "keywords":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            # Support both first_line_indent: true/720 and explicit values
            fi = element.get("first_line_indent", True)
            if fi is True or (isinstance(fi, (int, float)) and fi > 0):
                p.paragraph_format.first_line_indent = Inches(0.5)
            else:
                p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("label", "Keywords: "))
            run.italic = element.get("label_italic", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            # Support both field names: "items" (MD prompt) and "keywords" (legacy)
            keywords = element.get("items", element.get("keywords", []))
            if isinstance(keywords, list):
                keywords_text = ", ".join(str(k) for k in keywords)
            else:
                keywords_text = str(keywords)
            run2 = p.add_run(keywords_text)
            run2.font.name = 'Times New Roman'
            run2.font.size = Pt(12)


def _write_body(doc, section_data, instructions):
    """Write all body content (intro through discussion)."""
    doc.add_section()
    # Support both field names: body_first_line_indent_dxa (MD prompt) and body_first_line_indent (legacy)
    indent_dxa = instructions.get("body_first_line_indent_dxa", instructions.get("body_first_line_indent", 720))
    indent = Inches(indent_dxa / 1440) if indent_dxa > 13 else Inches(indent_dxa)

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "title_repeat":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", ""))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "heading":
            level = element.get("level", 1)
            level = min(max(level, 1), 3)
            p = doc.add_heading(element.get("text", ""), level=level)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)

        elif etype == "body_paragraph":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = indent
            _add_text_with_italics(p, element.get("text", ""))

        elif etype == "figure_caption":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0)
            label = element.get("label", f"Figure {element.get('number', '')}")
            run = p.add_run(label)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            caption = element.get("caption", "")
            if caption:
                p2 = doc.add_paragraph()
                p2.paragraph_format.line_spacing = 2.0
                p2.paragraph_format.first_line_indent = Inches(0)
                run2 = p2.add_run(caption)
                run2.italic = True
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(12)

        elif etype == "table_caption":
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.first_line_indent = Inches(0)
            label = element.get("label", f"Table {element.get('number', '')}")
            run = p.add_run(label)
            run.bold = True
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            caption = element.get("caption", "")
            if caption:
                p2 = doc.add_paragraph()
                p2.paragraph_format.line_spacing = 2.0
                p2.paragraph_format.first_line_indent = Inches(0)
                run2 = p2.add_run(caption)
                run2.italic = True
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(12)


def _write_references_page(doc, section_data, instructions):
    """Write the references page with hanging indent."""
    doc.add_section()

    for element in section_data.get("elements", []):
        etype = element.get("type", "")

        if etype == "references_label":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(element.get("text", "References"))
            run.bold = element.get("bold", True)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

        elif etype == "reference_entry":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.line_spacing = 2.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.first_line_indent = Inches(-0.5)
            _add_text_with_italics(p, element.get("text", ""))


def _add_text_with_italics(paragraph, text):
    """Parse *italic* markers and create appropriate runs."""
    if not text:
        return
    parts = re.split(r'(\*[^*]+\*)', text)
    for part in parts:
        if part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            run = paragraph.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


def _add_page_number_header(section):
    """Add right-aligned page number to section header."""
    header = section.header
    header.is_linked_to_previous = False

    if not header.paragraphs:
        p = header.add_paragraph()
    else:
        p = header.paragraphs[0]

    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)

    run = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar1)

    run2 = p.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' PAGE '
    run2._r.append(instrText)

    run3 = p.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run3._r.append(fldChar2)

    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════════════════════
# LEGACY GENERIC WRITER (for non-APA journals)
# ═══════════════════════════════════════════════════════════════════════════════

def write_formatted_docx(instructions: dict, output_path: str) -> str:
    """
    Write a formatted DOCX file from structured docx_instructions.

    Handles both:
    - New APA format (sections with type=title_page/abstract_page/body/references_page)
    - Legacy flat format (sections with type=title/heading/paragraph/reference)
    """
    if not instructions or not instructions.get("sections"):
        raise DocumentWriteError("docx_instructions must contain a non-empty 'sections' list.")

    sections = instructions.get("sections", [])

    # Detect if this is the new APA section-based format
    section_types = {s.get("type") for s in sections if isinstance(s, dict)}
    is_apa_format = bool(section_types & {"title_page", "abstract_page", "references_page"})

    if is_apa_format:
        return build_apa_docx(instructions, output_path)

    # Legacy flat format
    doc = Document()
    rules = instructions.get("rules", {}) or {}

    doc_rules = rules.get("document", {})
    font_name = doc_rules.get("font", "Times New Roman")
    font_size = _safe_int(doc_rules.get("font_size", 12), 12)
    line_spacing = _safe_float(doc_rules.get("line_spacing", 2.0), 2.0)
    margins = doc_rules.get("margins", {})

    _apply_document_defaults(doc, font_name, font_size, line_spacing)
    _set_document_margins(doc, margins)

    for section in sections:
        section_type = section.get("type", "paragraph")
        content = section.get("content", "")
        if content is None:
            content = ""

        try:
            if section_type == "title":
                _add_title(doc, content, font_name, font_size)
            elif section_type == "heading":
                level = _safe_int(section.get("level", 1), 1)
                heading_rules = rules.get("headings", {}).get(f"H{level}", {})
                _add_heading(doc, content, level, heading_rules, font_name, font_size)
            elif section_type == "abstract":
                abstract_rules = rules.get("abstract", {})
                _add_abstract(doc, content, abstract_rules, font_name, font_size, line_spacing)
            elif section_type == "reference":
                ref_rules = rules.get("references", {})
                _add_reference(doc, content, ref_rules, font_name, font_size)
            elif section_type == "figure_caption":
                fig_rules = rules.get("figures", {})
                _add_figure_caption(doc, content, fig_rules, font_name, font_size)
            elif section_type == "table_caption":
                tbl_rules = rules.get("tables", {})
                _add_table_caption(doc, content, tbl_rules, font_name, font_size)
            else:
                _add_paragraph(doc, content, font_name, font_size, line_spacing)
        except Exception as e:
            logger.warning("[DOCX] Failed to render section type=%s: %s", section_type, e)
            try:
                _add_paragraph(doc, content, font_name, font_size, line_spacing)
            except Exception:
                pass

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    logger.info("DOCX written | file=%s | sections=%d", out.name, len(sections))
    return str(out)


# ═══════════════════════════════════════════════════════════════════════════════
# LEGACY HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_document_defaults(doc: Document, font_name: str, font_size: int, line_spacing: float) -> None:
    try:
        style = doc.styles["Normal"]
        style.font.name = font_name
        style.font.size = Pt(font_size)
        pf = style.paragraph_format
        _apply_line_spacing(pf, line_spacing)
    except Exception as e:
        logger.warning("[DOCX] Could not set document defaults: %s", e)


def _set_document_margins(doc: Document, margins: dict) -> None:
    for section in doc.sections:
        try:
            top = _parse_measurement(margins.get("top", 1.0))
            bottom = _parse_measurement(margins.get("bottom", 1.0))
            left = _parse_measurement(margins.get("left", 1.0))
            right = _parse_measurement(margins.get("right", 1.0))
            section.top_margin = Inches(top)
            section.bottom_margin = Inches(bottom)
            section.left_margin = Inches(left)
            section.right_margin = Inches(right)
        except Exception as e:
            logger.warning("[DOCX] Could not set margins: %s", e)


def _parse_measurement(value) -> float:
    if isinstance(value, (int, float)):
        return float(value)
    if isinstance(value, str):
        value = value.strip().lower()
        try:
            if value.endswith("in"):
                return float(value[:-2])
            elif value.endswith("cm"):
                return float(value[:-2]) / 2.54
            elif value.endswith("mm"):
                return float(value[:-2]) / 25.4
            elif value.endswith("pt"):
                return float(value[:-2]) / 72.0
            else:
                return float(value)
        except ValueError:
            logger.warning("[DOCX] Cannot parse measurement '%s' — defaulting to 1.0in", value)
            return 1.0
    return 1.0


def _apply_line_spacing(pf, line_spacing: float) -> None:
    if line_spacing == 1.0:
        pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    elif line_spacing == 1.5:
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    elif line_spacing == 2.0:
        pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    else:
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(line_spacing * 12)


def _apply_font(run, font_name: str, font_size: int, bold: bool = False, italic: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.italic = italic


def _apply_case_transform(text: str, case_rule: str) -> str:
    if not case_rule or case_rule not in _CASE_OPTIONS:
        return text
    try:
        if case_rule == "UPPERCASE":
            return text.upper()
        elif case_rule == "lowercase":
            return text.lower()
        elif case_rule == "Sentence case":
            if not text:
                return text
            return text[0].upper() + text[1:].lower() if len(text) > 1 else text.upper()
        elif case_rule == "Title Case":
            _SMALL_WORDS = {
                "a", "an", "the", "and", "but", "or", "for", "nor",
                "on", "at", "to", "by", "in", "of", "up", "as", "is",
            }
            words = text.split()
            result = []
            for i, word in enumerate(words):
                if len(word) >= 2 and word.isupper():
                    result.append(word)
                elif "-" in word:
                    result.append("-".join(
                        part.capitalize() for part in word.split("-")
                    ))
                elif i == 0 or word.lower() not in _SMALL_WORDS:
                    result.append(word.capitalize())
                else:
                    result.append(word.lower())
            return " ".join(result)
    except Exception as e:
        logger.warning("[DOCX] Case transform failed for '%s': %s", case_rule, e)
    return text


def _add_paragraph(doc: Document, text: str, font_name: str, font_size: int, line_spacing: float) -> None:
    para = doc.add_paragraph()
    run = para.add_run(text)
    _apply_font(run, font_name, font_size)
    _apply_line_spacing(para.paragraph_format, line_spacing)


def _add_title(doc: Document, text: str, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _apply_font(run, font_name, font_size, bold=True)


def _add_heading(doc: Document, text: str, level: int, heading_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = heading_rules.get("bold", True)
    italic = heading_rules.get("italic", False)
    centered = heading_rules.get("centered", False)
    case = heading_rules.get("case", "Title Case")
    heading_font_size = _safe_int(heading_rules.get("font_size", font_size), font_size)
    text = _apply_case_transform(text, case)
    if centered:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    _apply_font(run, font_name, heading_font_size, bold=bold, italic=italic)


def _add_abstract(doc: Document, text: str, abstract_rules: dict, font_name: str, font_size: int, line_spacing: float) -> None:
    label = abstract_rules.get("label", "Abstract")
    label_bold = abstract_rules.get("label_bold", False)
    centered = abstract_rules.get("label_centered", True)

    label_para = doc.add_paragraph()
    if centered:
        label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label_run = label_para.add_run(label)
    _apply_font(label_run, font_name, font_size, bold=label_bold)

    body_para = doc.add_paragraph()
    body_run = body_para.add_run(text)
    _apply_font(body_run, font_name, font_size)
    _apply_line_spacing(body_para.paragraph_format, line_spacing)


def _add_reference(doc: Document, text: str, ref_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    hanging = ref_rules.get("hanging_indent", True)
    hanging_inches = _parse_measurement(ref_rules.get("indent_size", ref_rules.get("hanging_indent_inches", 0.5)))
    if hanging:
        para.paragraph_format.left_indent = Inches(hanging_inches)
        para.paragraph_format.first_line_indent = Inches(-hanging_inches)
    run = para.add_run(text)
    _apply_font(run, font_name, font_size)


def _add_figure_caption(doc: Document, text: str, fig_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = fig_rules.get("label_bold", True)
    italic = fig_rules.get("caption_italic", False)
    run = para.add_run(text)
    _apply_font(run, font_name, font_size - 1, bold=bold, italic=italic)


def _add_table_caption(doc: Document, text: str, tbl_rules: dict, font_name: str, font_size: int) -> None:
    para = doc.add_paragraph()
    bold = tbl_rules.get("label_bold", True)
    run = para.add_run(text)
    _apply_font(run, font_name, font_size - 1, bold=bold)


# ═══════════════════════════════════════════════════════════════════════════════
# IN-PLACE DOCX TRANSFORMATION (for uploaded DOCX files)
# ═══════════════════════════════════════════════════════════════════════════════

def _apply_heading_style(paragraph, fix: dict) -> None:
    bold = fix.get("bold", True)
    italic = fix.get("italic", False)
    centered = fix.get("centered", False)
    case = fix.get("case", "Title Case")
    text = paragraph.text
    text = _apply_case_transform(text, case)
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.bold = bold
        run.italic = italic
    if text != paragraph.text and paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""


def transform_docx_in_place(
    source_docx_path: str,
    transform_data: dict,
    rules: dict,
    output_path: str,
) -> str:
    """Apply journal formatting transformations to a source DOCX in-place."""
    try:
        doc = Document(source_docx_path)
    except Exception as e:
        raise DocumentWriteError(f"Cannot open source DOCX '{source_docx_path}': {e}") from e

    doc_rules      = rules.get("document", {})
    headings_rules = rules.get("headings", {})
    ref_rules      = rules.get("references", {})

    font_name       = doc_rules.get("font", "Times New Roman")
    font_size_pt    = _safe_int(doc_rules.get("font_size", 12), 12)
    line_spacing_v  = _safe_float(doc_rules.get("line_spacing", 2.0), 2.0)
    margins         = doc_rules.get("margins", {})

    try:
        normal = doc.styles["Normal"]
        normal.font.name = font_name
        normal.font.size = Pt(font_size_pt)
        _apply_line_spacing(normal.paragraph_format, line_spacing_v)
    except Exception as e:
        logger.warning("[DOCX_INPLACE] Could not set Normal style: %s", e)

    for sec in doc.sections:
        try:
            sec.top_margin    = Inches(_parse_measurement(margins.get("top",    1.0)))
            sec.bottom_margin = Inches(_parse_measurement(margins.get("bottom", 1.0)))
            sec.left_margin   = Inches(_parse_measurement(margins.get("left",   1.0)))
            sec.right_margin  = Inches(_parse_measurement(margins.get("right",  1.0)))
        except Exception as e:
            logger.warning("[DOCX_INPLACE] Could not set margins: %s", e)

    for para in doc.paragraphs:
        if not para.style.name.startswith("Heading"):
            continue
        try:
            level = int(para.style.name.split()[-1])
        except (ValueError, IndexError):
            level = 1

        h_rules  = headings_rules.get(f"H{level}", {})
        case     = h_rules.get("case", None)
        bold     = h_rules.get("bold", True)
        italic   = h_rules.get("italic", False)
        centered = h_rules.get("centered", False)
        h_size   = _safe_int(h_rules.get("font_size", font_size_pt), font_size_pt)

        if centered:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if case:
            new_text = _apply_case_transform(para.text, case)
            if new_text != para.text:
                _replace_paragraph_text_preserve_runs(para, new_text)

        for run in para.runs:
            try:
                run.bold        = bold
                run.italic      = italic
                run.font.name   = font_name
                run.font.size   = Pt(h_size)
            except Exception as e:
                logger.warning("[DOCX_INPLACE] Heading run format error: %s", e)

    reference_order = transform_data.get("reference_order", [])
    _reorder_references_in_place(doc, reference_order, ref_rules)

    out = Path(output_path).resolve()
    out.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(str(out))
    except Exception as e:
        raise DocumentWriteError(f"Failed to save in-place DOCX to '{out}': {e}") from e

    logger.info("[DOCX_INPLACE] Saved — file=%s", out.name)
    return str(out)


def _replace_paragraph_text_preserve_runs(para, new_text: str) -> None:
    if not para.runs:
        para.text = new_text
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def _reorder_references_in_place(doc: Document, reference_order: list, ref_rules: dict) -> None:
    paragraphs = doc.paragraphs
    ref_start_idx: Optional[int] = None
    for i, para in enumerate(paragraphs):
        if re.search(r"^\s*references?\s*$", para.text, re.IGNORECASE):
            ref_start_idx = i
            break

    if ref_start_idx is None:
        return

    ref_paras = []
    for para in paragraphs[ref_start_idx + 1:]:
        if para.style.name.startswith("Heading"):
            break
        if para.text.strip():
            ref_paras.append(para)

    if not ref_paras:
        return

    if reference_order and len(reference_order) == len(ref_paras):
        ordered = reference_order
    else:
        ordering = ref_rules.get("ordering", "alphabetical")
        if ordering == "alphabetical":
            ordered = sorted([p.text for p in ref_paras], key=str.lower)
        else:
            ordered = [p.text for p in ref_paras]

    for para, new_text in zip(ref_paras, ordered):
        if para.text != new_text:
            _replace_paragraph_text_preserve_runs(para, new_text)

    logger.info("[DOCX_INPLACE] References reordered — %d entries", len(ref_paras))


def _safe_int(value, default: int) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return default


def _safe_float(value, default: float) -> float:
    try:
        return float(value)
    except (TypeError, ValueError):
        return default
